"""
Resume text extraction from PDF and DOCX files.
Uses pdfplumber (better than PyPDF2 for structured text) and python-docx.
"""

import io
import re
from typing import Optional

import pdfplumber
from docx import Document


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file using pdfplumber."""
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file."""
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Dispatch to the right extractor based on file extension."""
    lower = filename.lower()
    if lower.endswith(".pdf"):
        text = extract_text_from_pdf(file_bytes)
    elif lower.endswith(".docx"):
        text = extract_text_from_docx(file_bytes)
    elif lower.endswith(".txt"):
        text = file_bytes.decode("utf-8", errors="ignore")
    else:
        raise ValueError(f"Unsupported file type: {filename}. Use PDF, DOCX, or TXT.")

    if not text.strip():
        raise ValueError("Could not extract any text from the uploaded file. The file may be image-based or corrupted.")

    return clean_text(text)


def clean_text(text: str) -> str:
    """Normalize whitespace and remove garbage characters."""
    # Remove null bytes and control chars (except newlines/tabs)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    # Collapse multiple blank lines into one
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Collapse multiple spaces
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


# ── Section detection ──────────────────────────────────────────────────────────

SECTION_PATTERNS = {
    "contact":     r"\b(contact|email|phone|address|linkedin|github|portfolio)\b",
    "summary":     r"\b(summary|objective|profile|about me|professional summary)\b",
    "experience":  r"\b(experience|work history|employment|career|positions? held)\b",
    "education":   r"\b(education|academic|degree|university|college|school|qualification)\b",
    "skills":      r"\b(skills?|technical skills?|competencies|expertise|technologies)\b",
    "projects":    r"\b(projects?|personal projects?|side projects?|portfolio)\b",
    "certifications": r"\b(certifications?|certificates?|licenses?|credentials?)\b",
    "achievements": r"\b(achievements?|awards?|honors?|accomplishments?)\b",
}


def detect_sections(text: str) -> dict[str, bool]:
    """Return which standard resume sections are present."""
    lower = text.lower()
    return {
        section: bool(re.search(pattern, lower))
        for section, pattern in SECTION_PATTERNS.items()
    }


def extract_contact_info(text: str) -> dict[str, Optional[str]]:
    """Pull email, phone, LinkedIn, GitHub from resume text."""
    email_match = re.search(
        r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}", text
    )
    phone_match = re.search(
        r"(\+?\d[\d\s\-().]{7,}\d)", text
    )
    linkedin_match = re.search(
        r"linkedin\.com/in/[\w\-]+", text, re.IGNORECASE
    )
    github_match = re.search(
        r"github\.com/[\w\-]+", text, re.IGNORECASE
    )
    return {
        "email":    email_match.group(0) if email_match else None,
        "phone":    phone_match.group(0).strip() if phone_match else None,
        "linkedin": linkedin_match.group(0) if linkedin_match else None,
        "github":   github_match.group(0) if github_match else None,
    }